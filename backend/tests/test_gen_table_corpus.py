import csv
import importlib.util
import io
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


PROJECT_ROOT = Path(__file__).resolve().parents[2]
SCRIPT_PATH = PROJECT_ROOT / "backend" / "scripts" / "gen_table_corpus.py"


def _load_generator_module():
    assert SCRIPT_PATH.is_file(), "table corpus generator script is missing"
    spec = importlib.util.spec_from_file_location("gen_table_corpus", SCRIPT_PATH)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def test_generate_table_corpus_creates_required_native_tables(tmp_path):
    module = _load_generator_module()

    generated = module.generate_table_corpus(tmp_path)

    assert {path.name for path in generated} == {
        "HTTP状态码速查表.docx",
        "Git常用命令对照表.xlsx",
        "Linux文本处理三剑客.csv",
    }

    document = Document(tmp_path / "HTTP状态码速查表.docx")
    assert len(document.tables) == 3
    assert [document.tables[0].rows[0].cells[index].text for index in range(3)] == [
        "状态码",
        "名称",
        "含义",
    ]
    assert all(len(table.rows) >= 11 for table in document.tables)

    workbook = load_workbook(tmp_path / "Git常用命令对照表.xlsx", read_only=True)
    assert workbook.sheetnames == ["基础操作", "分支与远程"]
    for worksheet in workbook.worksheets:
        assert [worksheet.cell(row=1, column=index).value for index in range(1, 4)] == [
            "命令",
            "作用",
            "示例",
        ]
        assert worksheet.max_row >= 13
    workbook.close()

    with zipfile.ZipFile(tmp_path / "Git常用命令对照表.xlsx", "r") as package:
        assert b"<autoFilter" not in package.read("xl/worksheets/sheet1.xml")
        assert b"<autoFilter" not in package.read("xl/worksheets/sheet2.xml")
        core_properties = package.read("docProps/core.xml")
        assert b"2026-08-23T00:00:00Z" in core_properties
        assert core_properties.count(b"2026-08-23T00:00:00Z") == 2

    with (tmp_path / "Linux文本处理三剑客.csv").open(
        "r", encoding="utf-8", newline=""
    ) as stream:
        rows = list(csv.reader(stream))
    assert rows[0] == ["工具", "用法", "示例", "说明"]
    assert len(rows) >= 16
    assert {row[0] for row in rows[1:]} == {"grep", "sed", "awk"}


def test_generate_table_corpus_is_byte_identical_on_rerun(tmp_path):
    module = _load_generator_module()

    first_paths = module.generate_table_corpus(tmp_path)
    first_bytes = {path.name: path.read_bytes() for path in first_paths}
    second_paths = module.generate_table_corpus(tmp_path)

    assert {path.name: path.read_bytes() for path in second_paths} == first_bytes


def test_normalize_office_zip_preserves_entry_order_and_metadata():
    module = _load_generator_module()
    source = io.BytesIO()
    with zipfile.ZipFile(source, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        first = zipfile.ZipInfo("z.xml", (2026, 8, 23, 12, 0, 0))
        first.compress_type = zipfile.ZIP_DEFLATED
        first.comment = b"keep-comment"
        first.external_attr = 0o600 << 16
        archive.writestr(first, b"z")
        archive.writestr("a.xml", b"a")

    normalized = module._normalize_office_zip(source.getvalue())

    with zipfile.ZipFile(io.BytesIO(source.getvalue()), "r") as before, zipfile.ZipFile(
        io.BytesIO(normalized), "r"
    ) as after:
        assert after.namelist() == before.namelist()
        assert after.getinfo("z.xml").comment == before.getinfo("z.xml").comment
        assert after.getinfo("z.xml").external_attr == before.getinfo("z.xml").external_attr
        assert all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in after.infolist())
