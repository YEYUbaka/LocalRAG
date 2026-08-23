"""Generate deterministic native-table corpus files for retrieval evaluation."""

from __future__ import annotations

import argparse
import copy
import csv
import io
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.table import Table, TableStyleInfo


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "test_docs"
FIXED_TIMESTAMP = datetime(2026, 8, 23, tzinfo=timezone.utc)
ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
CORE_TIMESTAMP = b"2026-08-23T00:00:00Z"


HTTP_TABLES = (
    (
        "2xx 成功",
        (
            ("200", "OK", "请求成功，响应中返回所请求的资源或结果。"),
            ("201", "Created", "请求成功并创建了新的资源。"),
            ("202", "Accepted", "请求已接受处理，但处理尚未完成。"),
            ("203", "Non-Authoritative Information", "返回的元信息来自本地或第三方副本，而非源服务器。"),
            ("204", "No Content", "请求成功，但响应没有消息正文。"),
            ("205", "Reset Content", "请求成功，客户端应重置当前文档视图。"),
            ("206", "Partial Content", "服务器按 Range 请求返回资源的一部分。"),
            ("207", "Multi-Status", "响应正文包含多个资源各自的状态。"),
            ("208", "Already Reported", "DAV 绑定成员已在响应前面列出，不再重复枚举。"),
            ("226", "IM Used", "服务器完成请求，并对当前实例应用了一个或多个实例操作。"),
        ),
    ),
    (
        "4xx 客户端错误",
        (
            ("400", "Bad Request", "请求语法、格式或参数无效，服务器无法处理。"),
            ("401", "Unauthorized", "请求缺少有效的身份认证凭据。"),
            ("402", "Payment Required", "为将来的数字支付场景保留的状态码。"),
            ("403", "Forbidden", "服务器理解请求，但拒绝授权访问。"),
            ("404", "Not Found", "服务器找不到请求的资源。"),
            ("405", "Method Not Allowed", "目标资源不支持请求所用的 HTTP 方法。"),
            ("406", "Not Acceptable", "服务器无法生成符合客户端内容协商条件的响应。"),
            ("407", "Proxy Authentication Required", "客户端必须先通过代理服务器认证。"),
            ("408", "Request Timeout", "服务器等待客户端请求的时间超过限制。"),
            ("409", "Conflict", "请求与资源当前状态发生冲突。"),
            ("410", "Gone", "目标资源已永久移除且没有转发地址。"),
            ("411", "Length Required", "服务器要求请求包含有效的 Content-Length。"),
        ),
    ),
    (
        "5xx 服务器错误",
        (
            ("500", "Internal Server Error", "服务器遇到未预期情况，无法完成请求。"),
            ("501", "Not Implemented", "服务器不支持完成请求所需的功能。"),
            ("502", "Bad Gateway", "网关或代理从上游服务器收到无效响应。"),
            ("503", "Service Unavailable", "服务器暂时无法处理请求，通常是过载或维护。"),
            ("504", "Gateway Timeout", "网关或代理未及时收到上游服务器响应。"),
            ("505", "HTTP Version Not Supported", "服务器不支持请求使用的 HTTP 版本。"),
            ("506", "Variant Also Negotiates", "服务器存在循环的透明内容协商配置。"),
            ("507", "Insufficient Storage", "服务器没有足够存储空间完成请求。"),
            ("508", "Loop Detected", "服务器处理请求时检测到无限循环。"),
            ("510", "Not Extended", "请求需要进一步扩展才能被服务器满足。"),
            ("511", "Network Authentication Required", "客户端需要完成网络访问认证。"),
        ),
    ),
)


GIT_SHEETS = (
    (
        "基础操作",
        (
            ("git init", "初始化新的 Git 仓库", "git init"),
            ("git clone", "克隆远程仓库", "git clone https://example.com/team/repo.git"),
            ("git status", "查看工作区和暂存区状态", "git status --short"),
            ("git add", "把变更加入暂存区", "git add README.md"),
            ("git commit", "提交暂存区中的变更", "git commit -m \"docs: update readme\""),
            ("git diff", "查看尚未暂存的差异", "git diff"),
            ("git diff --cached", "查看已暂存但未提交的差异", "git diff --cached"),
            ("git log", "查看提交历史", "git log --oneline --decorate"),
            ("git show", "查看某次提交的内容", "git show HEAD"),
            ("git restore", "恢复工作区文件", "git restore README.md"),
            ("git reset", "调整暂存区或移动当前分支指针", "git reset HEAD README.md"),
            ("git stash", "临时保存未提交的工作", "git stash push -m \"wip\""),
        ),
    ),
    (
        "分支与远程",
        (
            ("git branch", "列出或创建分支", "git branch feature/search"),
            ("git switch -c", "创建并切换到新分支", "git switch -c feature/search"),
            ("git switch", "切换到已有分支", "git switch main"),
            ("git merge", "把指定分支合并到当前分支", "git merge feature/search"),
            ("git rebase", "把当前分支提交变基到新基线", "git rebase origin/main"),
            ("git tag", "创建或列出标签", "git tag -a v1.0.0 -m \"release v1.0.0\""),
            ("git remote -v", "查看远程仓库地址", "git remote -v"),
            ("git remote add", "添加远程仓库", "git remote add origin https://example.com/team/repo.git"),
            ("git fetch", "下载远程引用但不合并", "git fetch origin"),
            ("git pull --rebase", "拉取远程提交并将本地提交变基", "git pull --rebase origin main"),
            ("git push -u", "首次推送并设置上游分支", "git push -u origin feature/search"),
            ("git push --delete", "删除远程分支", "git push origin --delete feature/search"),
        ),
    ),
)


LINUX_ROWS = (
    ("grep", "显示匹配行号", "grep -n 'error' app.log", "在匹配结果前输出源文件行号。"),
    ("grep", "递归搜索目录", "grep -r 'TODO' src/", "递归查找目录中包含指定文本的文件。"),
    ("grep", "忽略大小写", "grep -i 'warning' app.log", "匹配时不区分英文字母大小写。"),
    ("grep", "反向筛选", "grep -v '^#' config.ini", "只输出不匹配模式的行。"),
    ("grep", "使用扩展正则", "grep -E 'error|warning' app.log", "用扩展正则同时匹配多个模式。"),
    ("sed", "输出指定行范围", "sed -n '1,10p' file.txt", "关闭默认输出，仅打印第 1 到 10 行。"),
    ("sed", "全局替换文本", "sed 's/old/new/g' file.txt", "替换每行中出现的全部目标文本。"),
    ("sed", "原地修改并备份", "sed -i.bak 's/old/new/g' file.txt", "修改原文件并保留 .bak 备份。"),
    ("sed", "删除空行", "sed '/^$/d' file.txt", "从输出中移除空白内容为空的行。"),
    ("sed", "打印模式范围", "sed -n '/start/,/end/p' file.txt", "输出两个匹配模式之间的所有行。"),
    ("awk", "输出第一列", "awk '{print $1}' file.txt", "按默认空白分隔字段并打印第一列。"),
    ("awk", "指定字段分隔符", "awk -F, '{print $2}' data.csv", "使用逗号分隔字段并打印第二列。"),
    ("awk", "按数值条件筛选", "awk '$3 > 80 {print $1, $3}' scores.txt", "只输出第三列大于 80 的记录。"),
    ("awk", "保留表头并筛选", "awk 'NR==1 || $3 > 80' scores.txt", "输出第一行表头以及满足条件的记录。"),
    ("awk", "计算列总和", "awk '{sum += $1} END {print sum}' numbers.txt", "累计第一列并在输入结束后输出总和。"),
)


def _normalize_office_zip(payload: bytes) -> bytes:
    source = io.BytesIO(payload)
    target = io.BytesIO()
    with zipfile.ZipFile(source, "r") as reader, zipfile.ZipFile(
        target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as writer:
        for original in reader.infolist():
            info = copy.copy(original)
            info.date_time = ZIP_TIMESTAMP
            payload = reader.read(original.filename)
            if original.filename == "docProps/core.xml":
                for tag in (b"created", b"modified"):
                    payload = re.sub(
                        rb"(<dcterms:" + tag + rb"[^>]*>)[^<]*(</dcterms:" + tag + rb">)",
                        rb"\g<1>" + CORE_TIMESTAMP + rb"\g<2>",
                        payload,
                    )
            writer.writestr(info, payload)
    return target.getvalue()


def _write_if_changed(path: Path, payload: bytes) -> None:
    if path.is_file() and path.read_bytes() == payload:
        return
    path.write_bytes(payload)


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _build_http_docx() -> bytes:
    document = Document()
    document.core_properties.title = "HTTP 状态码速查表"
    document.core_properties.subject = "LocalRAG Golden Set 原生表格评测语料"
    document.core_properties.author = "LocalRAG"
    document.core_properties.created = FIXED_TIMESTAMP
    document.core_properties.modified = FIXED_TIMESTAMP
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = document.add_heading("HTTP 状态码速查表", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = document.add_paragraph(
        "按响应类别整理常见 HTTP 状态码。每个分类使用独立的原生 Word 表格。"
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for index, (heading, rows) in enumerate(HTTP_TABLES, start=1):
        if index > 1:
            document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(f"表 {index}：{heading}", level=1)
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.autofit = False
        widths = (Cm(2.2), Cm(6.0), Cm(9.0))
        for column, text in enumerate(("状态码", "名称", "含义")):
            cell = table.rows[0].cells[column]
            cell.text = text
            cell.width = widths[column]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _shade_cell(cell, "D9EAF7")
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10.5)
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for code, name, meaning in rows:
            cells = table.add_row().cells
            for column, text in enumerate((code, name, meaning)):
                cells[column].text = text
                cells[column].width = widths[column]
                cells[column].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cells[column].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9.5)

    stream = io.BytesIO()
    document.save(stream)
    return _normalize_office_zip(stream.getvalue())


def _build_git_xlsx() -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)
    workbook.properties.title = "Git 常用命令对照表"
    workbook.properties.subject = "LocalRAG Golden Set 原生表格评测语料"
    workbook.properties.creator = "LocalRAG"
    workbook.properties.created = FIXED_TIMESTAMP.replace(tzinfo=None)
    workbook.properties.modified = FIXED_TIMESTAMP.replace(tzinfo=None)

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    light_border = Border(bottom=Side(style="thin", color="B4C6E7"))

    for sheet_name, rows in GIT_SHEETS:
        worksheet = workbook.create_sheet(sheet_name)
        worksheet.sheet_view.showGridLines = False
        worksheet.freeze_panes = "A2"
        worksheet.append(("命令", "作用", "示例"))
        for row in rows:
            worksheet.append(row)
        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        for row in worksheet.iter_rows(min_row=2, max_row=worksheet.max_row):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = light_border
        worksheet.column_dimensions["A"].width = 24
        worksheet.column_dimensions["B"].width = 34
        worksheet.column_dimensions["C"].width = 52
        worksheet.row_dimensions[1].height = 24
        table_name = "GitBasicCommands" if sheet_name == "基础操作" else "GitBranchRemoteCommands"
        excel_table = Table(displayName=table_name, ref=f"A1:C{worksheet.max_row}")
        excel_table.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium2",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        worksheet.add_table(excel_table)

    stream = io.BytesIO()
    workbook.save(stream)
    workbook.close()
    return _normalize_office_zip(stream.getvalue())


def _build_linux_csv() -> bytes:
    stream = io.StringIO(newline="")
    writer = csv.writer(stream, lineterminator="\n")
    writer.writerow(("工具", "用法", "示例", "说明"))
    writer.writerows(LINUX_ROWS)
    return stream.getvalue().encode("utf-8")


def generate_table_corpus(output_dir: Path) -> tuple[Path, ...]:
    """Generate all native-table evaluation corpus files under ``output_dir``."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    artifacts = (
        (output_dir / "HTTP状态码速查表.docx", _build_http_docx()),
        (output_dir / "Git常用命令对照表.xlsx", _build_git_xlsx()),
        (output_dir / "Linux文本处理三剑客.csv", _build_linux_csv()),
    )
    for path, payload in artifacts:
        _write_if_changed(path, payload)
    return tuple(path for path, _ in artifacts)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=DEFAULT_OUTPUT_DIR,
        help=f"语料输出目录（默认：{DEFAULT_OUTPUT_DIR}）",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    for path in generate_table_corpus(args.output_dir):
        print(path.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
